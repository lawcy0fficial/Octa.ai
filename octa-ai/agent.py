"""
╔══════════════════════════════════════════════════════════════╗
║         NEXT GENERATION AI AGENT v1.0 — ULTIMATE            ║
║━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━║
║  Inspired by: Claude.ai                                      ║
║  Platform:    Parrot OS / Kali Linux                         ║
║  Cost:        $0/month forever                               ║
║                                                              ║
║  CAPABILITIES:                                               ║
║  ✅ Multi-provider rotation (82,500+ req/day free)           ║
║  ✅ Discussion → Planning → Execution → Review               ║
║  ✅ Agentic loop (never stops, auto-fixes errors)            ║
║  ✅ Claude Code + Ollama local (unlimited)                   ║
║  ✅ Open WebUI GUI (Claude.ai-like interface)                ║
║  ✅ MCP full Parrot OS access                                ║
║  ✅ Streaming responses (word by word)                       ║
║  ✅ Persistent memory (ChromaDB)                             ║
║  ✅ Vision (image analysis)                                  ║
║  ✅ Web search + CVE research                                ║
║  ✅ OSINT automation                                         ║
║  ✅ Security research tools                                   ║
║  ✅ PDF/MD/DOCX/HTML output                                  ║
║  ✅ Image generation (Pollinations.ai free)                  ║
║  ✅ Voice synthesis (Coqui TTS local)                        ║
║  ✅ Data analysis (Pandas + AI)                              ║
║  ✅ Live code execution                                       ║
╚══════════════════════════════════════════════════════════════╝
"""

OCTA_BANNER = """
╔══════════════════════════════════════════════════════════════╗
║                                                              ║
║    ██████╗  ██████╗████████╗ █████╗      █████╗ ██╗        ║
║   ██╔═══██╗██╔════╝╚══██╔══╝██╔══██╗    ██╔══██╗██║        ║
║   ██║   ██║██║        ██║   ███████║    ███████║██║        ║
║   ██║   ██║██║        ██║   ██╔══██║    ██╔══██║██║        ║
║   ╚██████╔╝╚██████╗   ██║   ██║  ██║    ██║  ██║██║        ║
║    ╚═════╝  ╚═════╝   ╚═╝   ╚═╝  ╚═╝    ╚═╝  ╚═╝╚═╝        ║
║                                                              ║
║         Omnipotent Cyber Threat Analysis AI v1.0            ║
║                                                              ║
║   🧠 Discussion  📋 Planning  ⚙️  Execution  ✅ Review       ║
║   🔐 Security    🔍 OSINT     🎨 Creative   📄 Reports       ║
║                                                              ║
║   Platform: Parrot OS / Kali Linux | Cost: $0/month         ║
║   APIs: 82,500+ req/day FREE | Local: ♾️  Unlimited         ║
║                                                              ║
╚══════════════════════════════════════════════════════════════╝
"""

import litellm
import os
import time
import ast
import json
import subprocess
import hashlib
import requests
from itertools import cycle
from datetime import datetime
from pathlib import Path
from typing import Optional, List, Dict, Tuple

# ══════════════════════════════════════════════════════════════
# CONFIGURATION
# ══════════════════════════════════════════════════════════════

litellm.set_verbose = False

# ── Load Keys from File or Env ────────────────────────────────
def load_keys() -> dict:
    """Load keys from keys.json — same search order as keys.py so
    this works no matter what directory agent.py is launched from."""
    locations = [
        Path(__file__).resolve().parent / "keys.json",
        Path("keys.json"),
        Path.home() / "agent" / "keys.json",
    ]
    for keys_file in locations:
        if keys_file.exists():
            try:
                with open(keys_file) as f:
                    return json.load(f)
            except Exception as e:
                print(f"   ⚠️  Could not load {keys_file}: {e}")
    print("   ⚠️  No keys.json found in any known location — "
          "every provider call will fail auth until you fix this.")
    return {}

LOADED_KEYS = load_keys()

# ══════════════════════════════════════════════════════════════
# ALL API KEYS — FULL ROTATION
# ══════════════════════════════════════════════════════════════

# NOTE: no hardcoded fallback keys — populate keys.json instead.
# Using empty/placeholder defaults keeps no live-looking secrets in source.
OPENROUTER_KEYS = cycle(LOADED_KEYS.get("openrouter", []) or [""])
CEREBRAS_KEYS   = cycle(LOADED_KEYS.get("cerebras", []) or [""])
GROK_KEYS       = cycle(LOADED_KEYS.get("grok", []) or [""])
GEMINI_KEYS     = cycle(LOADED_KEYS.get("gemini", []) or [""])
REQUESTY_KEYS   = cycle(LOADED_KEYS.get("requesty", []) or [""])
NOVITA_KEYS     = cycle(LOADED_KEYS.get("novita", []) or [""])

HUGGINGFACE_KEY  = (LOADED_KEYS.get("huggingface", []) or [""])[0]
POLLINATIONS_KEY = (LOADED_KEYS.get("pollinations", []) or [""])[0]
CLOUDFLARE_KEY   = (LOADED_KEYS.get("cloudflare", []) or [""])[0]

# ══════════════════════════════════════════════════════════════
# KEY ROTATION HELPERS
# ══════════════════════════════════════════════════════════════

def rotate_env_keys():
    """Rotate all provider keys in environment."""
    os.environ["OPENROUTER_API_KEY"] = next(OPENROUTER_KEYS)
    os.environ["CEREBRAS_API_KEY"]   = next(CEREBRAS_KEYS)
    os.environ["XAI_API_KEY"]        = next(GROK_KEYS)
    os.environ["GEMINI_API_KEY"]     = next(GEMINI_KEYS)
    os.environ["HF_TOKEN"]           = HUGGINGFACE_KEY

rotate_env_keys()

# ══════════════════════════════════════════════════════════════
# MODEL POOLS — Complete Priority System
# ══════════════════════════════════════════════════════════════
#
# OpenRouter entries are pulled live from keys.fetch_openrouter_free_models()
# instead of being hardcoded here. Several IDs that used to be hardcoded in
# this file no longer exist at all (there was never a free Claude 3 Haiku on
# OpenRouter, and DeepSeek/Mistral dropped their free variants) — those dead
# IDs were exactly why the agent's "auto-pick a free model" kept failing.
from keys import fetch_openrouter_free_models as _fetch_free_models
_LIVE_FREE_MODELS = _fetch_free_models(limit=6)
_OR_POOL_ENTRIES = [(f"openrouter/{m}", "openrouter") for m in _LIVE_FREE_MODELS]

# Best models for reasoning & discussion
DISCUSSION_POOL = cycle(_OR_POOL_ENTRIES + [
    ("cerebras/llama3.3-70b",                        "cerebras"),
    ("gemini/gemini-2.5-flash",                      "gemini"),
    ("xai/grok-beta",                                "grok"),
])

# Fastest models for planning
PLANNING_POOL = cycle([
    ("cerebras/llama3.3-70b",                        "cerebras"),
    ("gemini/gemini-2.5-flash",                      "gemini"),
    ("xai/grok-beta",                                "grok"),
] + _OR_POOL_ENTRIES)

# LOCAL ONLY — Unlimited, no filters, private
EXECUTION_POOL = cycle([
    ("ollama/phi3:mini",    "ollama"),
    ("ollama/qwen2.5:3b",   "ollama"),
    ("ollama/llama3.2:3b",  "ollama"),
])

# High quality models for review
REVIEW_POOL = cycle(_OR_POOL_ENTRIES + [
    ("gemini/gemini-2.5-flash",                      "gemini"),
    ("cerebras/llama3.3-70b",                        "cerebras"),
    ("xai/grok-beta",                                "grok"),
])

# Vision-capable models
VISION_POOL = cycle([
    ("gemini/gemini-2.5-flash",                      "gemini"),
    ("xai/grok-beta",                                "grok"),
])

# ══════════════════════════════════════════════════════════════
# CORE AI CALL ENGINE
# ══════════════════════════════════════════════════════════════

def call_ai(messages: list,
            pool,
            label: str = "",
            retries: int = 12,
            max_tokens: int = 4000) -> Tuple[str, str]:
    """
    Universal AI caller with full key + provider rotation.
    Auto-retries with different providers on failure.
    Returns (response_text, model_used).
    """
    for attempt in range(retries):
        model, provider = next(pool)

        # Rotate keys before each call
        rotate_env_keys()

        try:
            icon = "▶" if attempt == 0 else "🔁"
            print(f"   {icon} [{label}] {model[:50]}")

            kwargs: dict = dict(
                model=model,
                messages=messages,
                max_tokens=max_tokens,
                temperature=0.7,
                timeout=60
            )

            # Provider-specific config
            if provider == "ollama":
                kwargs["api_base"] = "http://localhost:11434"
            elif provider == "requesty":
                kwargs["api_base"] = "https://router.requesty.ai/v1"
                kwargs["api_key"]  = next(REQUESTY_KEYS)
            elif provider == "novita":
                kwargs["api_base"] = "https://api.novita.ai/v3/openai"
                kwargs["api_key"]  = next(NOVITA_KEYS)

            response = litellm.completion(**kwargs)
            result   = response.choices[0].message.content

            print(f"   ✅ [{label}] ← {model[:40]}")
            return result, model

        except Exception as e:
            err = str(e)[:80]
            if "429" in err or "rate" in err.lower():
                print(f"   ⏳ Rate limit → rotating key+provider...")
            elif "401" in err or "auth" in err.lower():
                print(f"   🔑 Auth error → rotating key...")
            elif "timeout" in err.lower():
                print(f"   ⌛ Timeout → switching provider...")
            else:
                print(f"   ⚠️  Error: {err[:50]}")
            time.sleep(1)

    return "ERROR: All providers exhausted.", "none"

def call_ai_stream(messages: list, label: str = ""):
    """Stream response word by word like Claude.ai."""
    model, provider = next(DISCUSSION_POOL)
    rotate_env_keys()

    try:
        kwargs = dict(
            model=model,
            messages=messages,
            stream=True,
            max_tokens=3000
        )
        if provider == "ollama":
            kwargs["api_base"] = "http://localhost:11434"

        response = litellm.completion(**kwargs)
        full     = ""
        for chunk in response:
            token = chunk.choices[0].delta.content or ""
            if token:
                print(token, end="", flush=True)
                full += token
        print()
        return full
    except Exception as e:
        print(f"\nStream error: {e}")
        result, _ = call_ai(messages, DISCUSSION_POOL, label)
        return result

# ══════════════════════════════════════════════════════════════
# MEMORY SYSTEM
# ══════════════════════════════════════════════════════════════

_memory_available = False
_conversations    = None
_research_mem     = None

def init_memory():
    """Initialize ChromaDB memory."""
    global _memory_available, _conversations, _research_mem
    try:
        import chromadb
        Path("memory").mkdir(exist_ok=True)
        client          = chromadb.PersistentClient(path="./memory")
        _conversations  = client.get_or_create_collection("conversations")
        _research_mem   = client.get_or_create_collection("research")
        _memory_available = True
        print("   🧠 Memory: ChromaDB initialized")
    except Exception as e:
        print(f"   ⚠️  Memory unavailable: {e}")

def save_to_memory(content: str, meta: dict = {}):
    """Save to persistent memory."""
    if not _memory_available or not _conversations:
        return
    try:
        doc_id = hashlib.md5(
            f"{content}{time.time()}".encode()
        ).hexdigest()
        _conversations.add(
            documents=[content],
            metadatas=[{"ts": datetime.now().isoformat(), **meta}],
            ids=[doc_id]
        )
    except Exception:
        pass

def get_memory_context(query: str, n: int = 3) -> str:
    """Retrieve relevant past context."""
    if not _memory_available or not _conversations:
        return ""
    try:
        res = _conversations.query(
            query_texts=[query], n_results=n)
        docs = res.get("documents", [[]])[0]
        if not docs:
            return ""
        context = "## Relevant Past Context:\n"
        for doc in docs:
            context += f"- {doc[:200]}\n"
        return context
    except Exception:
        return ""

# ══════════════════════════════════════════════════════════════
# SYSTEM PROMPTS
# ══════════════════════════════════════════════════════════════

PROMPTS = {
    "discuss": (
        "You are a world-class expert analyst and software architect. "
        "Deeply analyze the request. Identify: "
        "1) Core requirements and objectives "
        "2) Technical approach and best stack "
        "3) Edge cases and challenges "
        "4) Best implementation strategy. "
        "Be thorough, precise, and actionable."
    ),
    "plan": (
        "You are a senior software architect. "
        "Create a detailed, numbered implementation plan. "
        "Include: file structure, functions/classes, "
        "dependencies, error handling, testing approach. "
        "Make every step clear and actionable."
    ),
    "execute": (
        "You are an expert software engineer. "
        "Write COMPLETE, WORKING code. Rules: "
        "- NO placeholders or TODO comments "
        "- NO '...' gaps "
        "- FULL implementation only "
        "- Add error handling "
        "- Include ALL imports "
        "- Add comments "
        "- Code must run without modification "
        "- Add # TASK_COMPLETE when fully done"
    ),
    "review": (
        "You are a senior code reviewer. "
        "Review for: bugs, missing imports, security, "
        "performance, logic errors. "
        "Return COMPLETE fixed code only."
    ),
    "security_analysis": (
        "You are a senior security researcher. "
        "Analyze findings professionally. Provide: "
        "critical vulnerabilities, CVSS scores, "
        "risk assessment, remediation steps. "
        "Human reviewer will validate before any action."
    ),
    "osint_analysis": (
        "You are an expert OSINT analyst. "
        "Analyze intelligence data and provide: "
        "key findings, patterns, connections, "
        "security concerns, follow-up recommendations."
    ),
}

# ══════════════════════════════════════════════════════════════
# PHASE 1: DISCUSSION
# ══════════════════════════════════════════════════════════════

def discussion_phase(task: str) -> str:
    """Analyze and understand task using best online AI."""
    print("\n" + "═"*60)
    print("🗣️  PHASE 1 — DISCUSSION & ANALYSIS")
    print("   Providers: OpenRouter×10 + Cerebras×9 + Grok×9")
    print("              Gemini×7 + Novita×8 + Requesty×10")
    print("═"*60)

    # Get memory context
    context = get_memory_context(task)
    system  = PROMPTS["discuss"]
    if context:
        system += f"\n\n{context}"

    result, model = call_ai(
        [{"role": "system", "content": system},
         {"role": "user",   "content": task}],
        DISCUSSION_POOL, label="discuss"
    )
    print(f"\n📝 Analysis complete ← {model[:40]}")
    return result

# ══════════════════════════════════════════════════════════════
# PHASE 2: PLANNING
# ══════════════════════════════════════════════════════════════

def planning_phase(discussion: str) -> str:
    """Create implementation plan using fast online AI."""
    print("\n" + "═"*60)
    print("📋  PHASE 2 — PLANNING")
    print("   Providers: Cerebras×9 + Gemini×7 + Grok×9")
    print("═"*60)

    result, model = call_ai(
        [{"role": "system", "content": PROMPTS["plan"]},
         {"role": "user",
          "content": f"Create plan for:\n\n{discussion}"}],
        PLANNING_POOL, label="plan"
    )
    print(f"\n📋 Plan ready ← {model[:40]}")
    return result

# ══════════════════════════════════════════════════════════════
# PHASE 3: EXECUTION — AGENTIC LOOP (NEVER STOPS)
# ══════════════════════════════════════════════════════════════

def execution_phase(plan: str,
                    task: str,
                    max_iter: int = 30) -> str:
    """
    Execute with never-stops agentic loop.
    LOCAL OLLAMA ONLY — unlimited, private, no filters.
    Auto-saves every iteration.
    Auto-fixes syntax errors automatically.
    """
    print("\n" + "═"*60)
    print("⚙️   PHASE 3 — EXECUTION (LOCAL OLLAMA — Unlimited)")
    print("   Models: phi3:mini + qwen2.5:3b + llama3.2:3b")
    print("   No content filters | No rate limits | Private")
    print("═"*60)

    code = ""
    Path("outputs").mkdir(exist_ok=True)

    for i in range(1, max_iter + 1):
        print(f"\n   🔄 Iteration {i}/{max_iter}")

        # Generate/continue code
        new_code, model = call_ai(
            [{"role": "system", "content": PROMPTS["execute"]},
             {"role": "user",   "content":
                 f"Task: {task}\n\n"
                 f"Plan:\n{plan}\n\n"
                 f"Current code (iter {i}):\n"
                 f"{code if code else 'Start fresh.'}\n\n"
                 "Continue. Fix all issues. Never stop halfway. "
                 "Add # TASK_COMPLETE when 100% done."}],
            EXECUTION_POOL,
            label=f"exec-{i}",
            max_tokens=4000
        )
        code = new_code

        # Auto-save every iteration
        Path(f"outputs/v{i}.py").write_text(code, encoding="utf-8")
        Path("outputs/latest.py").write_text(code, encoding="utf-8")
        print(f"   💾 Auto-saved v{i}.py")

        # Auto syntax check + fix
        try:
            ast.parse(code)
            print("   ✅ Syntax clean!")

            # Check completion
            if "# TASK_COMPLETE" in code:
                print(f"\n   🎉 COMPLETE in {i} iterations!")
                Path("outputs/FINAL.py").write_text(
                    code, encoding="utf-8")
                break

            # Ask AI if done
            verify, _ = call_ai(
                [{"role": "system",
                  "content": "Answer ONLY: YES or NO"},
                 {"role": "user",
                  "content":
                      f"Is this code 100% complete for: '{task}'?\n\n"
                      f"{code[:1500]}"}],
                REVIEW_POOL, label="verify", max_tokens=10
            )
            if "YES" in verify.upper():
                print(f"\n   🎉 AI confirmed complete!")
                Path("outputs/FINAL.py").write_text(
                    code, encoding="utf-8")
                break
            else:
                print("   📝 Continuing...")

        except SyntaxError as e:
            print(f"   ⚠️  Syntax error: {e}")
            print("   🔧 Auto-fixing...")
            fixed, _ = call_ai(
                [{"role": "system",
                  "content":
                      "Fix this Python syntax error exactly. "
                      "Return ONLY complete fixed code."},
                 {"role": "user",
                  "content": f"Error: {e}\n\nCode:\n{code}"}],
                EXECUTION_POOL, label="fix"
            )
            code = fixed
            Path("outputs/latest.py").write_text(
                code, encoding="utf-8")
            print("   ✅ Fixed!")

        time.sleep(0.5)

    return code

# ══════════════════════════════════════════════════════════════
# PHASE 4: REVIEW
# ══════════════════════════════════════════════════════════════

def review_phase(code: str) -> str:
    """Final quality review using best online AI."""
    print("\n" + "═"*60)
    print("✅  PHASE 4 — REVIEW & POLISH")
    print("   Providers: All pools rotating")
    print("═"*60)

    result, model = call_ai(
        [{"role": "system", "content": PROMPTS["review"]},
         {"role": "user",
          "content": f"Review and fix:\n\n{code}"}],
        REVIEW_POOL, label="review", max_tokens=4000
    )
    print(f"\n✅ Review complete ← {model[:40]}")
    return result

# ══════════════════════════════════════════════════════════════
# OUTPUT GENERATION
# ══════════════════════════════════════════════════════════════

def save_all_formats(content: str,
                     task: str,
                     discussion: str = "",
                     plan: str = "") -> dict:
    """Save output in all formats: py, md, pdf, docx, html."""
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    Path("outputs").mkdir(exist_ok=True)
    saved = {}

    # Python
    py_path = f"outputs/{ts}.py"
    Path(py_path).write_text(content, encoding="utf-8")
    saved["py"] = py_path
    print(f"   💾 PY   → {py_path}")

    # Markdown
    md_path = f"outputs/{ts}.md"
    md = (f"# {task}\n\n"
          f"*{datetime.now().strftime('%Y-%m-%d %H:%M')}*\n\n"
          f"---\n\n"
          + (f"## Analysis\n{discussion}\n\n" if discussion else "")
          + (f"## Plan\n{plan}\n\n" if plan else "")
          + f"## Code\n```python\n{content}\n```")
    Path(md_path).write_text(md, encoding="utf-8")
    saved["md"] = md_path
    print(f"   💾 MD   → {md_path}")

    # PDF
    try:
        from fpdf import FPDF
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 10, task[:80], ln=True, align="C")
        pdf.set_font("Arial", "I", 9)
        pdf.cell(0, 6,
                 datetime.now().strftime('%Y-%m-%d %H:%M'),
                 ln=True, align="C")
        pdf.ln(3)
        pdf.set_font("Courier", size=8)
        for line in content.split("\n"):
            safe = line[:115].encode(
                'latin-1', 'replace').decode('latin-1')
            pdf.cell(0, 4, txt=safe, ln=True)
        pdf_path = f"outputs/{ts}.pdf"
        pdf.output(pdf_path)
        saved["pdf"] = pdf_path
        print(f"   💾 PDF  → {pdf_path}")
    except Exception as e:
        print(f"   ⚠️  PDF: {e}")

    # DOCX
    try:
        from docx import Document
        doc = Document()
        doc.add_heading(task[:100], 0)
        doc.add_paragraph(
            datetime.now().strftime('%Y-%m-%d %H:%M'))
        if discussion:
            doc.add_heading("Analysis", level=1)
            doc.add_paragraph(discussion)
        if plan:
            doc.add_heading("Plan", level=1)
            doc.add_paragraph(plan)
        doc.add_heading("Code", level=1)
        doc.add_paragraph(content)
        docx_path = f"outputs/{ts}.docx"
        doc.save(docx_path)
        saved["docx"] = docx_path
        print(f"   💾 DOCX → {docx_path}")
    except Exception as e:
        print(f"   ⚠️  DOCX: {e}")

    # HTML
    html_path = f"outputs/{ts}.html"
    html = f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8">
<title>{task[:60]}</title>
<style>
body{{background:#0d1117;color:#c9d1d9;font-family:monospace;
     max-width:1000px;margin:0 auto;padding:20px}}
h1{{color:#58a6ff;border-bottom:2px solid #21262d;padding-bottom:10px}}
h2{{color:#3fb950}}
pre{{background:#161b22;padding:20px;border-radius:8px;
    border-left:4px solid #58a6ff;overflow-x:auto;
    white-space:pre-wrap;word-wrap:break-word}}
.meta{{color:#8b949e;font-size:0.9em}}
</style></head>
<body>
<h1>🤖 {task[:80]}</h1>
<p class="meta">Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
<h2>Code</h2>
<pre>{content.replace('<','&lt;').replace('>','&gt;')}</pre>
</body></html>"""
    Path(html_path).write_text(html, encoding="utf-8")
    saved["html"] = html_path
    print(f"   💾 HTML → {html_path}")

    print(f"\n✅ All saved with base: outputs/{ts}")
    return saved

# ══════════════════════════════════════════════════════════════
# VISION — IMAGE ANALYSIS
# ══════════════════════════════════════════════════════════════

def analyze_image(image_path: str,
                  question: str = "Describe this image in detail.") -> str:
    """Analyze image using Gemini free vision."""
    import base64
    try:
        with open(image_path, "rb") as f:
            img_data = base64.b64encode(f.read()).decode()
        ext       = Path(image_path).suffix.lower()
        mime_type = {"jpg":"image/jpeg","jpeg":"image/jpeg",
                     "png":"image/png","gif":"image/gif",
                     "webp":"image/webp"}.get(ext[1:], "image/jpeg")

        os.environ["GEMINI_API_KEY"] = next(GEMINI_KEYS)
        response = litellm.completion(
            model="gemini/gemini-2.5-flash",
            messages=[{"role": "user", "content": [
                {"type": "image_url",
                 "image_url": {
                     "url": f"data:{mime_type};base64,{img_data}"}},
                {"type": "text", "text": question}
            ]}],
            max_tokens=2000
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Vision error: {e}"

# ══════════════════════════════════════════════════════════════
# WEB SEARCH
# ══════════════════════════════════════════════════════════════

def web_search(query: str, n: int = 5) -> str:
    """Search web using DuckDuckGo (free, no key)."""
    try:
        from duckduckgo_search import DDGS
        results = []
        with DDGS() as ddgs:
            for r in ddgs.text(query, max_results=n):
                results.append(
                    f"**{r.get('title')}**\n"
                    f"{r.get('href')}\n"
                    f"{r.get('body', '')[:200]}\n"
                )
        return "\n".join(results) if results else "No results"
    except Exception as e:
        return f"Search error: {e}"

def search_cves(service: str, version: str = "") -> list:
    """Search NVD for CVEs (free, no auth)."""
    try:
        r = requests.get(
            "https://services.nvd.nist.gov/rest/json/cves/2.0"
            f"?keywordSearch={service}+{version}&resultsPerPage=10",
            timeout=15)
        cves = []
        for v in r.json().get("vulnerabilities", []):
            cve   = v.get("cve", {})
            score = (cve.get("metrics", {})
                    .get("cvssMetricV31", [{}])[0]
                    .get("cvssData", {})
                    .get("baseScore", "N/A"))
            cves.append({
                "id":    cve.get("id"),
                "score": score,
                "desc":  cve.get("descriptions",
                                 [{}])[0].get("value", "")[:150]
            })
        return sorted(cves,
                      key=lambda x: float(x.get("score") or 0),
                      reverse=True)
    except Exception as e:
        return [{"error": str(e)}]

# ══════════════════════════════════════════════════════════════
# OSINT TOOLS
# ══════════════════════════════════════════════════════════════

def run_tool(command: str, timeout: int = 120) -> str:
    """Run system tool safely."""
    try:
        result = subprocess.run(
            command, shell=True,
            capture_output=True, text=True,
            timeout=timeout)
        return result.stdout or result.stderr or "No output"
    except subprocess.TimeoutExpired:
        return f"Timeout after {timeout}s"
    except Exception as e:
        return f"Error: {e}"

def ensure_tool(tool: str) -> bool:
    """Check tool installed, install if missing."""
    import shutil
    if shutil.which(tool):
        return True
    print(f"   📦 Installing {tool}...")
    r = subprocess.run(
        f"sudo apt install -y {tool}",
        shell=True, capture_output=True)
    return r.returncode == 0

def domain_osint(domain: str) -> dict:
    """Quick domain OSINT."""
    results = {}
    results['whois']   = run_tool(f"whois {domain}", 30)
    results['dns']     = run_tool(
        f"dig {domain} ANY +noall +answer")
    ensure_tool("subfinder")
    results['subs']    = run_tool(
        f"subfinder -d {domain} -silent", 60)
    ensure_tool("whatweb")
    results['tech']    = run_tool(f"whatweb {domain}", 30)
    return results

# ══════════════════════════════════════════════════════════════
# IMAGE GENERATION
# ══════════════════════════════════════════════════════════════

def generate_image(prompt: str,
                   filename: str = None) -> str:
    """Generate image via Pollinations.ai (100% free)."""
    import urllib.parse
    if not filename:
        ts       = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"outputs/image_{ts}.jpg"
    Path("outputs").mkdir(exist_ok=True)
    encoded = urllib.parse.quote(prompt)
    url     = (f"https://image.pollinations.ai/prompt/{encoded}"
               f"?width=1024&height=1024&nologo=true")
    try:
        r = requests.get(url, timeout=60)
        if r.status_code == 200:
            Path(filename).write_bytes(r.content)
            return f"✅ Saved: {filename}"
        return f"Error: HTTP {r.status_code}"
    except Exception as e:
        return f"Error: {e}"

# ══════════════════════════════════════════════════════════════
# VOICE SYNTHESIS
# ══════════════════════════════════════════════════════════════

def text_to_speech(text: str,
                   output: str = None) -> str:
    """Convert text to speech (Coqui TTS, free local)."""
    if not output:
        ts     = datetime.now().strftime("%Y%m%d_%H%M%S")
        output = f"outputs/speech_{ts}.wav"
    Path("outputs").mkdir(exist_ok=True)
    result = subprocess.run(
        f'tts --text "{text[:500]}" --out_path {output}',
        shell=True, capture_output=True, text=True)
    if result.returncode == 0:
        return f"✅ Speech: {output}"
    return f"TTS error: {result.stderr[:100]}"

# ══════════════════════════════════════════════════════════════
# CLAUDE CODE INTEGRATION
# ══════════════════════════════════════════════════════════════

def setup_claude_code(model: str = None):
    """
    Configure Claude Code.

    NOTE: this used to route Claude Code through OpenRouter with
    ANTHROPIC_BASE_URL + a free Llama/DeepSeek model pretending to be
    Claude. That can't work — Claude Code speaks the Anthropic-native
    Messages API, not OpenRouter's OpenAI-compatible format, and
    OpenRouter has never had a free Claude model. Real Claude access
    is either a Claude subscription (`claude` login) or a real
    console.anthropic.com key in keys.json under "anthropic". This
    now delegates to claude_code_setup.py, which handles both.
    """
    try:
        from claude_code_setup import setup_claude_code as _real_setup
        return _real_setup()
    except Exception as e:
        print(f"⚠️  Could not run claude_code_setup.py: {e}")
        print("   Run it directly: python3 claude_code_setup.py")
        return None

# ══════════════════════════════════════════════════════════════
# MAIN AGENT RUNNER
# ══════════════════════════════════════════════════════════════

def run_agent(task: str) -> str:
    """
    Run complete 4-phase agent for any task.
    Discussion → Planning → Execution → Review
    """
    print(f"\n{'╔' + '═'*58 + '╗'}")
    print(f"║{'🚀 AGENT v1.0 — STARTED':^58}║")
    print(f"{'╚' + '═'*58 + '╝'}")
    print(f"\n📥 Task: {task}")

    start = time.time()

    # Save task to memory
    save_to_memory(f"Task: {task}", {"type": "task"})

    # Run all phases
    discussion = discussion_phase(task)
    plan       = planning_phase(discussion)
    code       = execution_phase(plan, task)
    final      = review_phase(code)

    # Save all formats
    print("\n💾 Saving outputs...")
    saved = save_all_formats(final, task, discussion, plan)

    # Save result to memory
    save_to_memory(
        f"Task: {task}\nResult: {final[:500]}",
        {"type": "result"}
    )

    elapsed = round(time.time() - start, 2)
    print(f"\n{'╔' + '═'*58 + '╗'}")
    print(f"║{'🎯 COMPLETED in ' + str(elapsed) + 's':^58}║")
    print(f"{'╚' + '═'*58 + '╝'}")

    # Preview
    preview = final[:600] + "\n..." if len(final) > 600 else final
    print(f"\n📄 Preview:\n{preview}")

    return final

# ══════════════════════════════════════════════════════════════
# CAPACITY DISPLAY
# ══════════════════════════════════════════════════════════════

def show_capacity():
    print("""
╔══════════════════════════════════════════════════════════════╗
║              FREE API CAPACITY — v1.0                        ║
╠══════════════════════════════════════════════════════════════╣
║  Provider      Keys  Daily Limit          Speed              ║
║  ──────────────────────────────────────────────────────────  ║
║  OpenRouter    ×10   200×28models=56,000  Varies            ║
║  Cerebras      ×9    ~1M tokens/key=9M    ⚡ 2,000 tok/s    ║
║  Grok/xAI      ×9    Large capacity       Fast              ║
║  Gemini        ×7    1,500×7=10,500/day   Fast (1M ctx)    ║
║  Requesty      ×10   200×8models=16,000   Fast              ║
║  Novita        ×8    Free tier            Fast              ║
║  HuggingFace   ×1    2,000/day            Moderate          ║
║  Ollama        ♾️    UNLIMITED local       RAM dependent     ║
╠══════════════════════════════════════════════════════════════╣
║  TOTAL CLOUD:  ~82,500+ requests/day FREE                    ║
║  LOCAL:        ♾️  Unlimited via Ollama                     ║
║  COST:         $0/month forever                              ║
╚══════════════════════════════════════════════════════════════╝
""")

# ══════════════════════════════════════════════════════════════
# ENTRY POINT
# ══════════════════════════════════════════════════════════════

if __name__ == "__main__":
    # Initialize memory
    init_memory()
    show_capacity()

    print("""
╔══════════════════════════════════════════════════════════════╗
║          OCTA AI v1.0 — ULTIMATE                   ║
╠══════════════════════════════════════════════════════════════╣
║  Commands:                                                    ║
║  [any task]  → Run 4-phase agent                            ║
║  stream      → Stream response word by word                  ║
║  image       → Generate image (Pollinations.ai)              ║
║  vision      → Analyze an image                              ║
║  search      → Web search                                     ║
║  cve         → CVE research                                  ║
║  osint       → OSINT on domain                               ║
║  setup-cc    → Configure Claude Code                         ║
║  capacity    → Show API capacity                             ║
║  quit        → Exit                                          ║
╚══════════════════════════════════════════════════════════════╝
""")

    while True:
        try:
            user_input = input("\n💬 > ").strip()
            if not user_input:
                continue

            if user_input.lower() in ["quit", "exit", "q"]:
                print("👋 Goodbye!")
                break

            elif user_input.lower() == "capacity":
                show_capacity()

            elif user_input.lower() == "setup-cc":
                setup_claude_code()

            elif user_input.lower().startswith("stream "):
                task = user_input[7:].strip()
                print(f"\n📡 Streaming: {task}\n")
                call_ai_stream(
                    [{"role": "user", "content": task}],
                    label="stream")

            elif user_input.lower().startswith("image "):
                prompt = user_input[6:].strip()
                print(generate_image(prompt))

            elif user_input.lower().startswith("vision "):
                parts = user_input[7:].split(" ", 1)
                path  = parts[0]
                q     = parts[1] if len(parts) > 1 else "Describe this."
                print(analyze_image(path, q))

            elif user_input.lower().startswith("search "):
                query = user_input[7:].strip()
                print(web_search(query))

            elif user_input.lower().startswith("cve "):
                parts   = user_input[4:].strip().split()
                service = parts[0]
                version = parts[1] if len(parts) > 1 else ""
                results = search_cves(service, version)
                for c in results[:5]:
                    print(f"  {c.get('id')} "
                          f"(CVSS: {c.get('score')}) "
                          f"— {c.get('desc', '')[:80]}")

            elif user_input.lower().startswith("osint "):
                domain  = user_input[6:].strip()
                results = domain_osint(domain)
                result, _ = call_ai(
                    [{"role": "system",
                      "content": PROMPTS["osint_analysis"]},
                     {"role": "user",
                      "content": json.dumps(results)[:5000]}],
                    DISCUSSION_POOL, label="osint")
                print(result)

            elif user_input.lower().startswith("tts "):
                text = user_input[4:].strip()
                print(text_to_speech(text))

            else:
                run_agent(user_input)

        except KeyboardInterrupt:
            print("\n\n👋 Interrupted!")
            break
        except Exception as e:
            print(f"\n❌ Error: {e}")
            print("   Retrying...")
