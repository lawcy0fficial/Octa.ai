"""
Output Generation v1.0 — Fully Connected
==========================================
PDF | MD | DOCX | HTML | JSON reports
Imported by: ALL modules for consistent reporting
"""

import json
from datetime import datetime
from pathlib import Path

# ══════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════

def _ts() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")

def _ensure(*dirs):
    for d in dirs:
        Path(d).mkdir(parents=True, exist_ok=True)

# ══════════════════════════════════════════════════════════════
# MARKDOWN
# ══════════════════════════════════════════════════════════════

def save_markdown(content: str,
                  title: str = "Report",
                  filename: str = None) -> str:
    """Save as Markdown file."""
    _ensure("outputs")
    if not filename:
        filename = f"outputs/{_ts()}.md"

    md = (
        f"# {title}\n\n"
        f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n\n"
        f"---\n\n{content}"
    )
    Path(filename).write_text(md, encoding="utf-8")
    print(f"   💾 MD   → {filename}")
    return filename

# ══════════════════════════════════════════════════════════════
# PDF
# ══════════════════════════════════════════════════════════════

def save_pdf(content: str,
             title: str = "Report",
             filename: str = None) -> str:
    """Save as PDF with professional formatting."""
    try:
        from fpdf import FPDF

        _ensure("outputs")
        if not filename:
            filename = f"outputs/{_ts()}.pdf"

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        # Header
        pdf.set_fill_color(26, 26, 46)
        pdf.rect(0, 0, 210, 30, "F")
        pdf.set_text_color(255, 255, 255)
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 15, "", ln=True)
        pdf.cell(0, 10, title[:80], ln=True, align="C")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(5)

        # Date
        pdf.set_font("Arial", "I", 9)
        pdf.set_text_color(100, 100, 100)
        pdf.cell(0, 6,
                 f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
                 ln=True, align="C")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(3)

        # Content
        pdf.set_font("Courier", size=8)
        for line in content.split("\n"):
            try:
                # Handle headers
                if line.startswith("# "):
                    pdf.set_font("Arial", "B", 13)
                    pdf.set_text_color(0, 80, 200)
                    safe = line[2:].encode(
                        "latin-1", "replace").decode("latin-1")
                    pdf.cell(0, 8, safe[:100], ln=True)
                    pdf.set_font("Courier", size=8)
                    pdf.set_text_color(0, 0, 0)
                elif line.startswith("## "):
                    pdf.set_font("Arial", "B", 11)
                    pdf.set_text_color(0, 120, 60)
                    safe = line[3:].encode(
                        "latin-1", "replace").decode("latin-1")
                    pdf.cell(0, 7, safe[:100], ln=True)
                    pdf.set_font("Courier", size=8)
                    pdf.set_text_color(0, 0, 0)
                else:
                    safe = line[:115].encode(
                        "latin-1", "replace").decode("latin-1")
                    pdf.cell(0, 4, txt=safe, ln=True)
            except Exception:
                pass

        # Footer
        pdf.set_y(-15)
        pdf.set_font("Arial", "I", 8)
        pdf.set_text_color(150, 150, 150)
        pdf.cell(0, 10,
                 f"OCTA AI v1.0 | {datetime.now().strftime('%Y-%m-%d')}",
                 align="C")

        pdf.output(filename)
        print(f"   💾 PDF  → {filename}")
        return filename

    except ImportError:
        import subprocess
        subprocess.run(
            "pip3 install fpdf2 --break-system-packages",
            shell=True, capture_output=True)
        return save_pdf(content, title, filename)
    except Exception as e:
        print(f"   ⚠️  PDF error: {e}")
        return f"PDF error: {e}"

# ══════════════════════════════════════════════════════════════
# WORD DOCX
# ══════════════════════════════════════════════════════════════

def save_docx(content: str,
              title: str = "Report",
              filename: str = None,
              metadata: dict = {}) -> str:
    """Save as Word document."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        _ensure("outputs")
        if not filename:
            filename = f"outputs/{_ts()}.docx"

        doc = Document()

        # Title
        h   = doc.add_heading(title[:100], 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        meta_text = (
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        if metadata:
            for k, v in metadata.items():
                meta_text += f" | {k}: {v}"

        meta_para = doc.add_paragraph(meta_text)
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # Content — parse markdown headers
        for line in content.split("\n"):
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("```"):
                pass  # Skip code fences
            elif line.strip():
                doc.add_paragraph(line)

        doc.save(filename)
        print(f"   💾 DOCX → {filename}")
        return filename

    except ImportError:
        import subprocess
        subprocess.run(
            "pip3 install python-docx --break-system-packages",
            shell=True, capture_output=True)
        return save_docx(content, title, filename, metadata)
    except Exception as e:
        print(f"   ⚠️  DOCX error: {e}")
        return f"DOCX error: {e}"

# ══════════════════════════════════════════════════════════════
# HTML
# ══════════════════════════════════════════════════════════════

def save_html(content: str,
              title: str = "Report",
              filename: str = None) -> str:
    """Save as styled dark-theme HTML."""
    _ensure("outputs")
    if not filename:
        filename = f"outputs/{_ts()}.html"

    # Convert basic markdown to HTML
    html_content = ""
    in_code      = False
    for line in content.split("\n"):
        if line.startswith("```"):
            if in_code:
                html_content += "</pre>\n"
                in_code = False
            else:
                html_content += "<pre><code>\n"
                in_code = True
        elif in_code:
            html_content += (
                line.replace("<", "&lt;").replace(">", "&gt;")
                + "\n")
        elif line.startswith("# "):
            html_content += f"<h1>{line[2:]}</h1>\n"
        elif line.startswith("## "):
            html_content += f"<h2>{line[3:]}</h2>\n"
        elif line.startswith("### "):
            html_content += f"<h3>{line[4:]}</h3>\n"
        elif line.startswith("**") and line.endswith("**"):
            html_content += f"<strong>{line[2:-2]}</strong><br>\n"
        elif line.strip():
            html_content += f"<p>{line}</p>\n"

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width">
<title>{title}</title>
<style>
  :root {{
    --bg: #0d1117; --surface: #161b22;
    --border: #21262d; --text: #c9d1d9;
    --blue: #58a6ff; --green: #3fb950;
    --orange: #d29922; --red: #f85149;
  }}
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{
    background: var(--bg); color: var(--text);
    font-family: -apple-system, BlinkMacSystemFont,
                 'Segoe UI', sans-serif;
    max-width: 1000px; margin: 0 auto; padding: 20px;
    line-height: 1.6;
  }}
  .header {{
    background: linear-gradient(135deg, #1f2937, #111827);
    border: 1px solid var(--border);
    border-radius: 12px; padding: 30px;
    text-align: center; margin-bottom: 30px;
  }}
  .header h1 {{
    color: var(--blue); font-size: 2em;
    border: none; margin-bottom: 8px;
  }}
  .meta {{
    color: #8b949e; font-size: 0.85em;
  }}
  h1 {{
    color: var(--blue);
    border-bottom: 2px solid var(--border);
    padding-bottom: 8px; margin: 25px 0 15px;
  }}
  h2 {{ color: var(--green); margin: 20px 0 10px; }}
  h3 {{ color: var(--orange); margin: 15px 0 8px; }}
  p {{ margin: 8px 0; }}
  pre, code {{
    background: var(--surface);
    border: 1px solid var(--border);
    border-left: 4px solid var(--blue);
    border-radius: 6px; padding: 15px;
    overflow-x: auto; font-family: 'Courier New', monospace;
    font-size: 0.85em; white-space: pre-wrap;
    word-wrap: break-word; margin: 10px 0;
  }}
  strong {{ color: var(--blue); }}
  table {{
    border-collapse: collapse; width: 100%; margin: 15px 0;
  }}
  th, td {{
    border: 1px solid var(--border);
    padding: 8px 12px; text-align: left;
  }}
  th {{ background: var(--surface); color: var(--blue); }}
  .footer {{
    text-align: center; color: #8b949e;
    font-size: 0.8em; margin-top: 40px;
    padding-top: 20px;
    border-top: 1px solid var(--border);
  }}
</style>
</head>
<body>
<div class="header">
  <h1>🤖 {title}</h1>
  <p class="meta">
    Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
    | OCTA AI v1.0
  </p>
</div>
<div class="content">
{html_content}
</div>
<div class="footer">
  Generated by OCTA AI v1.0 | $0/month free
</div>
</body>
</html>"""

    Path(filename).write_text(html, encoding="utf-8")
    print(f"   💾 HTML → {filename}")
    return filename

# ══════════════════════════════════════════════════════════════
# JSON
# ══════════════════════════════════════════════════════════════

def save_json(data: dict,
              filename: str = None) -> str:
    """Save data as JSON file."""
    _ensure("outputs")
    if not filename:
        filename = f"outputs/{_ts()}.json"
    with open(filename, "w") as f:
        json.dump(data, f, indent=2, default=str)
    print(f"   💾 JSON → {filename}")
    return filename

# ══════════════════════════════════════════════════════════════
# SAVE ALL FORMATS
# ══════════════════════════════════════════════════════════════

def save_all_formats(content: str,
                     title: str = "Output",
                     base_dir: str = "outputs",
                     metadata: dict = {}) -> dict:
    """
    Save content in ALL formats simultaneously.
    Returns dict of saved file paths.
    """
    _ensure(base_dir)
    ts      = _ts()
    base    = f"{base_dir}/{ts}"
    saved   = {}

    print(f"\n💾 Saving all formats...")

    # Python (if it looks like code)
    if any(kw in content for kw in
           ["def ", "import ", "class ", "print("]):
        py_path = f"{base}.py"
        Path(py_path).write_text(content, encoding="utf-8")
        saved["py"] = py_path
        print(f"   💾 PY   → {py_path}")

    saved["md"]   = save_markdown(content, title, f"{base}.md")
    saved["pdf"]  = save_pdf(content, title, f"{base}.pdf")
    saved["docx"] = save_docx(content, title, f"{base}.docx",
                               metadata)
    saved["html"] = save_html(content, title, f"{base}.html")

    print(f"\n✅ Saved {len(saved)} formats → {base}.*")
    return saved

# ══════════════════════════════════════════════════════════════
# SECURITY REPORT TEMPLATE
# ══════════════════════════════════════════════════════════════

def security_report(target: str,
                    report_type: str,
                    findings: list,
                    analysis: str,
                    auth: str = "Authorized") -> str:
    """
    Generate professional security report.
    Saves as MD + PDF automatically.
    """
    ts   = _ts()
    _ensure("reports")

    report = f"""# {report_type.upper()} Security Report

**Target:** {target}
**Date:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
**Authorization:** {auth}
**Classification:** CONFIDENTIAL

---

## Executive Summary

{analysis[:600]}

---

## Findings ({len(findings)} total)

"""
    for i, f in enumerate(findings, 1):
        if isinstance(f, dict):
            report += (f"### Finding {i}: "
                       f"{f.get('type', 'Unknown')}\n"
                       f"**Severity:** {f.get('severity', 'N/A')}\n"
                       f"**Details:** {f.get('description', '')}\n\n")
        else:
            report += f"### Finding {i}\n{str(f)[:300]}\n\n"

    report += f"""---

## Full Analysis

{analysis}

---

## Disclaimer

*This report is for authorized security research only.*
*All testing performed with written authorization.*
*Generated by AI Security Agent v1.0*
"""

    # Save both MD and PDF
    report_path = f"reports/{report_type}_{target.replace('.','_')}_{ts}"
    md_path     = save_markdown(report, f"{report_type} Report", f"{report_path}.md")
    pdf_path    = save_pdf(report, f"{report_type} Report — {target}", f"{report_path}.pdf")

    print(f"\n📄 Reports saved:")
    print(f"   MD:  {md_path}")
    print(f"   PDF: {pdf_path}")
    return md_path

# ══════════════════════════════════════════════════════════════
# TEST
# ══════════════════════════════════════════════════════════════

if __name__ == "__main__":
    test_content = """# Test Report

## Summary
This is a test of the output generation module.

## Findings
- Finding 1: Everything working correctly
- Finding 2: All formats generating properly

## Code Example
```python
print("Hello from OCTA AI v1.0!")
```

## Conclusion
All output formats working correctly.
"""
    results = save_all_formats(test_content, "Test Report")
    print(f"\nSaved files:")
    for fmt, path in results.items():
        print(f"  {fmt}: {path}")
